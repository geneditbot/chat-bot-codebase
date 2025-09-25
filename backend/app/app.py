from fastapi import FastAPI, Request, Form, UploadFile, File, Query
from fastapi.middleware.cors import CORSMiddleware
from openai import OpenAI
import os
from dotenv import load_dotenv
from uuid import uuid4
from fastapi.responses import JSONResponse
from io import BytesIO
from docx import Document
import fitz  # PyMuPDF
from typing import Optional
from .init_db import init_db
from .models import ChatSession, Message, Feedback
from .db import SessionLocal
from fastapi.responses import StreamingResponse
from docx import Document
import io

#Load OpenAI API key
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

#Initialize DB
init_db()

#Initialize app
app = FastAPI()

# CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  # Vite default
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

#Define and initialize chatbot system prompt
SYSTEM_PROMPT = (
    '''You are an EDI advisor chatbot. Your role is to support educators in integrating Equity, Diversity, and Inclusion (EDI) principles into their ICT lesson plans. Draw on your knowledge of EDI in ICT education to offer thoughtful, practical, and constructive guidance.

Begin the conversation by warmly introducing yourself as an EDI advisor. Invite the educator to upload the lesson plan they’d like to enhance with EDI principles.

Only after the lesson plan is uploaded, ask the educator what kind of support they need. Present the following numbered options and prompt them to enter the number that best matches their needs. Let the educator know that support isn’t limited to these options—they may choose option 6 for any other type of assistance.

Support options (shown only after upload):

1. I want to integrate EDI principles into this lesson plan.

2. I want to include better examples or datasets that reflect EDI principles.

3. I want to design an EDI-integrated assignment for this lesson.

4. I want to include reflective questions to help students think about EDI in this lesson.

5. I want to evaluate my lesson plan in terms of how well it addresses EDI principles.

6. Something else.

If the educator selects a numbered option, respond with relevant insights, suggestions, or resources tailored to their choice. If they select “Something else,” ask them to describe their specific needs or goals.

 When offering suggestions, apply the following guiding principles:
 1. Strong Equity
Provide suggestions with a focus on strong equity, including:
• Recognition: Validate the lived experiences and knowledge of marginalized groups.
• Representation: Ensure students from diverse backgrounds are visible in content, examples, and discourse.
• Reframing: Challenge deficit narratives and stereotypes using inclusive language and critical reflection.

 2. Universal Design for Learning (UDL)
Apply UDL principles, especially those supporting emotional capacity:
• Embed empathy and restorative practices into learning activities.
• Use strategies that foster perspective-taking, relational awareness, and community trust.
• Design tasks that allow for multiple formats of expression and support safe academic risk-taking.

 3. Social Constructivist Learning
Promote collaborative learning and distributed expertise:
• Encourage peer interaction and co-construction of knowledge.
• Include content that raises awareness of different social groups to challenge assumptions.
• Use open-ended tasks that invite diverse perspectives and lived experiences.

 4. Teacher and Institutional Practice Awareness
Be mindful of hidden curriculum and institutional norms:
• Include diverse representation in texts, examples, and references.
• Avoid reinforcing dominant cultural norms or stereotypes.
• Design activities that disrupt bias and foster critical empathy.

 Design Requirements
• Offer multiple modes of engagement (e.g., visual, oral, written, experiential).
• Provide flexibility in how students demonstrate understanding.
• Use inclusive language and prompts that invite varied viewpoints.
• Include feedback mechanisms that are empathetic, growth-oriented, and restorative.
Where appropriate, integrate data or insights about different social groups to build awareness and counter deficit thinking.

Conversation Flow and Follow-up Guidance
Throughout the conversation:

Use a supportive, conversational tone.

Guide the educator with questions or prompts appropriate to their context.

Offer explanations, examples, or ideas suited to their level of experience with EDI.

If the educator seems unsure or stuck, suggest possible directions or ask clarifying questions.

If they enter an unrecognized input, gently prompt them to choose from the available options or rephrase their request.

Follow-up After Suggestions
After suggesting new content—such as examples, datasets, assignments, reflective questions, or learning activities—ask context-appropriate follow-up questions that help the educator reflect, refine, or move forward. These follow-up prompts should:

Encourage adaptation, integration, or deeper thinking;

Support decision-making about incorporating the suggestion;

Align with the educator’s original intent and lesson context;

Be supportive and conversational in tone.

Ask whether they would like to “update the lesson plan” definitely, whenever you provided content that can be directly added to the lesson plan. When asking this question please use the exact phrasing “update the lesson plan” within your message. 

At any point do not limit yourself only to the specifically mentioned follow-up question; 
including that question, include other relevant follow-up questions as well, according to the provided instructions.

Special Handling
If the educator chooses Option 2 (datasets/examples):

If only suggestions for improvement are offered, follow up by asking:
“Would you like to craft a sample dataset that reflects these principles?”

Only ask to "update the lesson plan" if a dataset or specific content has been generated.

If the educator chooses Option 4 (reflective questions):

After providing suggestions, ask whether they’d like to design an individual or group activity based on those questions.
'''
        )

#Initialize maximum number of messages in the chat history
MAX_HISTORY = 20

#Functions to extract text from the uploaded lesson plan
def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(BytesIO(file_bytes))
    return "\n".join(para.text for para in doc.paragraphs)

def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
        for page in pdf:
            text += page.get_text()
    return text

# Summarize older messages if available

def summarize_old_messages(session_id: str, db):
    all_messages = db.query(Message).filter_by(session_id=session_id).order_by(Message.timestamp).all()
    if len(all_messages) <= MAX_HISTORY:
        return None
    early_messages = all_messages[:-MAX_HISTORY]
    prompt = [
        {"role": "system", "content": "You are summarizing a conversation between an educator and an EDI advisor. Provide a brief summary of the conversation so far."},
        *[{"role": m.role, "content": m.content} for m in early_messages]
    ]
    summary_response = client.chat.completions.create(model="gpt-4.1-mini", messages=prompt)
    return summary_response.choices[0].message.content

#Send chat history
def get_chat_history(session_id: str, db):
    session = db.query(ChatSession).filter_by(id=session_id).first()
    history = db.query(Message).filter_by(session_id=session_id).order_by(Message.timestamp).all()
    messages = [{"role": m.role, "content": m.content} for m in history][-MAX_HISTORY:]

    system_prompt_present = any(
        m["role"] == "system" and SYSTEM_PROMPT in m["content"]
        for m in messages
    )

    #Inject system prompt if missing
    if not system_prompt_present:
        messages.insert(0, {
            "role": "system",
            "content": SYSTEM_PROMPT  
    })

    # Check if original lesson is referenced
    lesson_present = any(
        session.original_lesson.strip()[:100] in m["content"]
        for m in messages
    ) if session and session.original_lesson else False

    if session and session.original_lesson and not lesson_present:
        messages.insert(1, {
            "role": "user",
            "content": f"The original lesson plan for this conversation is:\n{session.original_lesson}"
        })

    #Inject chat history summary if available
    if session and session.summary:
        messages.insert(1, {
            "role": "system",
            "content": f"Summary of earlier conversation: {session.summary}"
        })

    return messages

@app.post("/chatStart")
async def chatStart():
    db = SessionLocal()
    session_id = str(uuid4())  # Create unique session ID

    # Initiate conversation executing system prompt
    db.add(ChatSession(id=session_id))
    db.add(Message(session_id=session_id, role="system", content=SYSTEM_PROMPT))

    response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages= [{"role": "system", "content": SYSTEM_PROMPT}]
        )
        
    api_response = response.choices[0].message.content

    db.add(Message(session_id=session_id, role="assistant", content=api_response))
    db.commit()
    db.close()

    return {"response": api_response, "session_id" :session_id}

@app.post("/chatContinue")
async def chatContinue(message: str = Form(...), session_id: str = Form(...)):

    db = SessionLocal()
    chat_session = db.query(ChatSession).filter_by(id=session_id).first()

    #Generate chat history summary if it is not available 
    if not chat_session.summary:
        summary = summarize_old_messages(session_id, db)
        if summary:
            chat_session.summary = summary #Set chat history summary to chat session
            db.commit()

    #Inject chat history and user message to the prompt
    chat_messages = get_chat_history(session_id, db)
    chat_messages.append({"role": "user", "content": message})

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=chat_messages
    )
    api_response = response.choices[0].message.content

    db.add(Message(session_id=session_id, role="user", content=message))
    db.add(Message(session_id=session_id, role="assistant", content=api_response))

    db.commit()
    db.close()

    return {"response": api_response, "session_id": session_id}

@app.post("/fileUpload")
async def chatStart(file: UploadFile = File(None), session_id: Optional[str] = Form(None)):
    db = SessionLocal()
    file_content=""
    #Extract content of the lesson plan
    if file:
        file_bytes = await file.read()
        if file.filename.endswith(".docx"):
            file_content = extract_text_from_docx(file_bytes)
        elif file.filename.endswith(".pdf"):
            file_content = extract_text_from_pdf(file_bytes)
        else:
            try:
                file_content = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                file_content = "[File uploaded, but not readable.]"

        #Retrieve chat session
        chat_session = db.query(ChatSession).filter_by(id=session_id).first()
        if not chat_session:
            return JSONResponse(status_code=400, content={"error": "Invalid session_id"})

        if not chat_session.summary:
            summary = summarize_old_messages(session_id, db)
            if summary:
                chat_session.summary = summary
                db.commit()

        #Set chat history and file content to the prompt
        chat_messages = get_chat_history(session_id, db)
        chat_messages.append({"role": "user", "content": f"Lesson Plan:\n{file_content}"})

        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=chat_messages
        )

        api_response = response.choices[0].message.content

        db.add(Message(session_id=session_id, role="user", content=f"Lesson Plan:\n{file_content}"))
        db.add(Message(session_id=session_id, role="assistant", content=api_response))
        chat_session.original_lesson = file_content #Update lesson plan in db with uploaded file content
        chat_session.updated_lesson = file_content
        db.commit()

        db.close()
        return {"response": api_response, "session_id": session_id}

#Retrieve chat sessions for chat history   
@app.get("/sessions")
def get_sessions():
        db = SessionLocal()
        sessions = db.query(ChatSession).filter(ChatSession.original_lesson.isnot(None)).order_by(ChatSession.created_at.desc()).all()
        results = []
        for s in sessions:
            results.append({
                "id": s.id,
                "created_at": s.created_at.isoformat() if s.created_at else None,
                "summary": s.summary if s.summary else "",
                "lesson_preview": (s.original_lesson[:100] + "...") if s.original_lesson else "",
            })
        db.close()
        return JSONResponse(content=results)

#Retrieve messages of the selected chat session from the chat history
@app.get("/sessionMessages")
def get_session_messages(session_id: str = Query(...)):
    file = ""
    db = SessionLocal()
    messages = db.query(Message).filter_by(session_id=session_id).order_by(Message.timestamp).all()
    session = db.query(ChatSession).filter_by(id=session_id).first()
    if session:
        if session.updated_lesson:
            file = session.updated_lesson
        else:
            file = session.original_lesson
            
    results = [{"role": m.role, "content": m.content} for m in messages]
    db.close()
    return {"file": file, "messages": results}

#Lesson plan update functionality
@app.post("/updateLesson")
async def update_lesson(session_id: str = Form(...), new_content: str = Form(...)):
    db = SessionLocal()
    currentContent =""
    chat_session = db.query(ChatSession).filter_by(id=session_id).first()
    if chat_session:
        if chat_session.updated_lesson:
            currentContent = chat_session.updated_lesson
        else:
            currentContent = chat_session.original_lesson
    if not chat_session.summary:
        summary = summarize_old_messages(session_id, db)
        if summary:
            chat_session.summary = summary
            db.commit()

    #Update lesson plan by appending suggested content using LLM API
    chat_messages = get_chat_history(session_id, db)
    chat_messages.append({"role": "user", "content": f"Update the lesson plan by integrating the new content - \n{new_content} in to the current lesson plan - \n{currentContent} appropriately preserving the pedagogical flow. In the response provide the full content of the updated lesson plan. At the start display Updated Lesson Plan as the heading. Do not include any additional texts in the response."})
    response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=chat_messages
        )

    api_response = response.choices[0].message.content
    db.add(Message(session_id=session_id, role="assistant", content=api_response))
    chat_session.updated_lesson = api_response #Update updated lesson plan in db 
    db.commit()
    db.close()
    return {"response": api_response, "session_id": session_id}

#Download updated lesson plan functionality
@app.get("/downloadLesson")
def download_lesson(session_id: str = Query(...)):
    db = SessionLocal()
    chat_session = db.query(ChatSession).filter_by(id=session_id).first()
    db.close()

    if not chat_session or not chat_session.updated_lesson:
        return JSONResponse(status_code=404, content={"error": "Updated lesson not found."})

    # Create a .docx document
    doc = Document()
    doc.add_heading("Updated Lesson Plan", level=1)
    for line in chat_session.updated_lesson.split("\n"):
        doc.add_paragraph(line)

    # Save to in-memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=updated_lesson_{session_id[:8]}.docx"
        }
    )

@app.post("/submitFeedback")
async def submit_feedback(session_id: str = Form(...), feedback: str = Form(...), feedbackProvider: str = Form(...)):
    db = SessionLocal()
    db.add(Feedback(session_id=session_id, feedback=feedback, name=feedbackProvider))
    db.commit()
    db.close()
    return {"message": "Feedback submitted successfully."}